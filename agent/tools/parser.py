"""Resume and job description parsing."""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any, Literal

if TYPE_CHECKING:
    from agent.tools.local_parser import ResumeStructured

import pdfplumber
from docx import Document

from agent.config import get_settings

FileFormat = Literal["pdf", "docx", "txt"]
RoleCategory = Literal[
    "software_engineering",
    "aiml",
    "data_science",
    "design",
    "research_academic",
    "non_portfolio",
]

_MUST_HAVE_PATTERNS = re.compile(
    r"(?im)^(?:[-*•]\s*)?(?:required|must\s+have|must-have|minimum|essential)\s*[:\-]?\s*(.+)$"
)
_NICE_TO_HAVE_PATTERNS = re.compile(
    r"(?im)^(?:[-*•]\s*)?(?:preferred|nice\s+to\s+have|nice-to-have|bonus|plus)\s*[:\-]?\s*(.+)$"
)
_BULLET = re.compile(r"(?m)^[-*•]\s+(.+)$")

_DOMAIN_KEYWORDS: dict[str, tuple[str, ...]] = {
    "technical": (
        "software",
        "engineer",
        "developer",
        "python",
        "backend",
        "frontend",
        "devops",
        "machine learning",
        "data scientist",
    ),
    "design": (
        "designer",
        "ux",
        "ui",
        "graphic",
        "figma",
        "creative",
        "visual",
    ),
    "academic": (
        "research",
        "professor",
        "phd",
        "postdoc",
        "university",
        "publication",
    ),
    "writing": (
        "writer",
        "content",
        "editor",
        "copywriter",
        "journalist",
    ),
    "business": (
        "product manager",
        "consultant",
        "startup",
        "founder",
        "sales",
    ),
}


@dataclass
class ParsedDocument:
    text: str
    hyperlinks: list[str] = field(default_factory=list)
    format: FileFormat = "txt"


VALID_REQUIREMENT_TYPES = frozenset(
    {
        "technical_skill",
        "soft_skill",
        "experience",
        "education",
        "responsibility",
    }
)


@dataclass
class JdRequirement:
    """Single JD requirement with optional LLM-assigned type."""

    text: str
    weight: Literal["must_have", "nice_to_have"] = "must_have"
    requirement_type: str | None = None


@dataclass
class JdStructured:
    job_title: str | None = None
    domain: str = "general"
    industry: str | None = None
    seniority: str | None = None
    role_category: RoleCategory = "non_portfolio"
    must_have: list[str] = field(default_factory=list)
    nice_to_have: list[str] = field(default_factory=list)
    requirements: list[JdRequirement] = field(default_factory=list)


def detect_format(content: bytes, filename: str = "") -> FileFormat:
    if content[:4] == b"%PDF":
        return "pdf"
    if content[:2] == b"PK" and filename.lower().endswith(".docx"):
        return "docx"
    if filename.lower().endswith(".pdf"):
        return "pdf"
    if filename.lower().endswith(".docx"):
        return "docx"
    return "txt"


def parse_file(content: bytes, filename: str = "") -> ParsedDocument:
    """Parse resume or JD bytes into plain text (and PDF hyperlinks)."""
    fmt = detect_format(content, filename)
    if fmt == "pdf":
        return _parse_pdf(content)
    if fmt == "docx":
        return _parse_docx(content)
    from agent.tools.local_parser import normalize_extracted_text

    return ParsedDocument(
        text=normalize_extracted_text(_decode_text(content)),
        hyperlinks=[],
        format="txt",
    )


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


def _parse_pdf(content: bytes) -> ParsedDocument:
    from agent.tools.local_parser import normalize_extracted_text

    text_parts: list[str] = []
    links: list[str] = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
            for link in page.hyperlinks or []:
                uri = (link.get("uri") or link.get("url") or "").strip()
                if uri:
                    links.append(uri)

    return ParsedDocument(
        text=normalize_extracted_text("\n\n".join(text_parts)),
        hyperlinks=links,
        format="pdf",
    )


def _parse_docx(content: bytes) -> ParsedDocument:
    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    from agent.tools.local_parser import normalize_extracted_text

    return ParsedDocument(
        text=normalize_extracted_text("\n".join(parts)),
        hyperlinks=[],
        format="docx",
    )


def _detect_domain(text: str) -> str:
    lowered = text.lower()
    scores = {
        domain: sum(1 for kw in keywords if kw in lowered)
        for domain, keywords in _DOMAIN_KEYWORDS.items()
    }
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "general"


def _detect_seniority(text: str) -> str | None:
    lowered = text.lower()
    for level in ("principal", "staff", "senior", "lead", "junior", "intern"):
        if re.search(rf"\b{re.escape(level)}\b", lowered):
            return level
    return None


def _extract_title(jd_text: str) -> str | None:
    for line in jd_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if len(stripped) < 120 and not stripped.startswith(("-", "*", "•")):
            return stripped
    return None


def _dedupe_lines(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            out.append(item)
    return out


def _extract_requirement_lines(jd_text: str) -> tuple[list[str], list[str]]:
    must: list[str] = []
    nice: list[str] = []

    for line in jd_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        for pattern in _MUST_HAVE_PATTERNS.finditer(stripped):
            must.append(pattern.group(1).strip())
        for pattern in _NICE_TO_HAVE_PATTERNS.finditer(stripped):
            nice.append(pattern.group(1).strip())

    in_must_section = False
    for line in jd_text.splitlines():
        lowered = line.lower().strip()
        if any(h in lowered for h in ("requirements", "must have", "qualifications")):
            in_must_section = True
            continue
        if in_must_section and any(h in lowered for h in ("nice to have", "preferred", "bonus")):
            in_must_section = False
        bullet = _BULLET.match(line.strip())
        if in_must_section and bullet:
            must.append(bullet.group(1).strip())

    in_nice_section = False
    for line in jd_text.splitlines():
        lowered = line.lower().strip()
        if any(h in lowered for h in ("preferred", "nice to have", "bonus")):
            in_nice_section = True
            continue
        bullet = _BULLET.match(line.strip())
        if in_nice_section and bullet:
            nice.append(bullet.group(1).strip())

    return _dedupe_lines(must), _dedupe_lines(nice)


def parse_jd_structured(jd_text: str, *, use_llm: bool | None = None) -> JdStructured:
    """
    Extract structured JD fields for rubric building.

    Uses Gemini when ``JD_PARSE_USE_LLM=true``, ``GEMINI_API_KEY`` is set, and
    ``use_llm`` is not False; falls back to heuristics on failure or when disabled.
    """
    if use_llm is None:
        settings = get_settings()
        from agent.llm_client import gemini_configured

        use_llm = settings.jd_parse_use_llm and gemini_configured(settings)

    if use_llm:
        try:
            return _parse_jd_with_gemini(jd_text)
        except Exception:
            pass

    return _parse_jd_heuristic(jd_text)


def _requirements_from_lists(
    must_have: list[str],
    nice_to_have: list[str],
) -> list[JdRequirement]:
    items: list[JdRequirement] = []
    for text in must_have:
        if text.strip():
            items.append(JdRequirement(text=text.strip(), weight="must_have"))
    for text in nice_to_have:
        if text.strip():
            items.append(JdRequirement(text=text.strip(), weight="nice_to_have"))
    return items


def _parse_requirements_from_gemini(data: dict[str, Any]) -> list[JdRequirement]:
    """Parse Gemini ``requirements`` array; fall back to string lists if absent."""
    raw_items = data.get("requirements")
    if isinstance(raw_items, list) and raw_items:
        parsed: list[JdRequirement] = []
        for item in raw_items:
            if not isinstance(item, dict):
                continue
            text = str(item.get("text") or item.get("requirement") or "").strip()
            if not text:
                continue
            weight = item.get("weight") or "must_have"
            if weight not in ("must_have", "nice_to_have"):
                weight = "must_have"
            req_type = item.get("requirement_type")
            if req_type not in VALID_REQUIREMENT_TYPES:
                req_type = None
            parsed.append(JdRequirement(text=text, weight=weight, requirement_type=req_type))
        if parsed:
            return parsed

    must = [str(x).strip() for x in (data.get("must_have") or []) if str(x).strip()]
    nice = [str(x).strip() for x in (data.get("nice_to_have") or []) if str(x).strip()]
    return _requirements_from_lists(must, nice)


def _parse_jd_heuristic(jd_text: str) -> JdStructured:
    from agent.tools.local_parser import parse_jd_local

    return parse_jd_local(jd_text)


def parse_resume_structured(text: str) -> ResumeStructured:
    """Extract resume highlights locally (no LLM)."""
    from agent.tools.local_parser import parse_resume_local

    return parse_resume_local(text)


def _parse_jd_with_gemini(jd_text: str) -> JdStructured:
    from google.genai import types

    from agent.llm_client import create_genai_client

    settings = get_settings()
    client = create_genai_client(settings)

    prompt = f"""Extract job description structure from the text below.
Return JSON only with keys:
job_title (string|null),
domain (technical|design|academic|writing|business|healthcare|general),
industry (string|null),
seniority (string|null),
role_category (one of: software_engineering, aiml, data_science, design,
research_academic, non_portfolio),
requirements (array of objects, each with:
  text (string),
  weight ("must_have" or "nice_to_have"),
  requirement_type (one of: technical_skill, soft_skill, experience, education, responsibility)
).

Classify role_category from title and duties (not brittle keyword rules):
- software_engineering: general SWE/backend/frontend/devops/platform roles
- aiml: ML/AI/LLM/deep-learning engineering or research engineering
- data_science: data scientist/analyst/BI/analytics roles
- design: UX/UI/product/graphic/visual design roles
- research_academic: academic/research/postdoc/publication-heavy roles
- non_portfolio: HR, sales, PM, operations, and other roles without standard portfolios

Classify requirement_type by meaning for ANY industry (e.g. RN license -> education,
years of experience -> experience, patient communication -> soft_skill).
Also include must_have and nice_to_have (string arrays) as duplicates of
requirements text.

JOB DESCRIPTION:
{jd_text[:12000]}
"""

    response = client.models.generate_content(
        model=settings.gemini_model_id,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.1,
        ),
    )
    raw = (response.text or "").strip()
    data: dict[str, Any] = json.loads(raw)

    requirements = _parse_requirements_from_gemini(data)
    must_have = [r.text for r in requirements if r.weight == "must_have"] or list(
        data.get("must_have") or []
    )
    nice_to_have = [r.text for r in requirements if r.weight == "nice_to_have"] or list(
        data.get("nice_to_have") or []
    )

    from agent.tools.portfolio_signal import infer_role_category, normalize_role_category

    role_category = normalize_role_category(data.get("role_category"))
    if role_category == "non_portfolio" and not data.get("role_category"):
        role_category = infer_role_category(
            job_title=data.get("job_title"),
            domain=str(data.get("domain") or "general"),
            jd_text=jd_text,
            must_have=must_have,
            nice_to_have=nice_to_have,
        )

    return JdStructured(
        job_title=data.get("job_title"),
        domain=str(data.get("domain") or "general"),
        industry=data.get("industry"),
        seniority=data.get("seniority"),
        role_category=role_category,
        must_have=must_have,
        nice_to_have=nice_to_have,
        requirements=requirements,
    )
