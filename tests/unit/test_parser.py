import io
from pathlib import Path

import pytest
from docx import Document

from agent.tools.parser import (
    detect_format,
    parse_file,
    parse_jd_structured,
)

FIXTURES = Path(__file__).resolve().parent.parent / "fixtures"


def test_detect_format_pdf_magic() -> None:
    assert detect_format(b"%PDF-1.4", "resume.pdf") == "pdf"


def test_detect_format_docx_magic() -> None:
    assert detect_format(b"PK\x03\x04", "resume.docx") == "docx"


def test_parse_plain_text_file() -> None:
    content = (FIXTURES / "sample_resume.txt").read_bytes()
    doc = parse_file(content, "resume.txt")
    assert "Senior Software Engineer" in doc.text
    assert doc.format == "txt"
    assert doc.hyperlinks == []


def test_parse_docx_generated() -> None:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("Backend engineer with Go and Rust.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table cell skill: gRPC"
    document.save(buffer)

    doc = parse_file(buffer.getvalue(), "resume.docx")
    assert "Backend engineer" in doc.text
    assert "gRPC" in doc.text
    assert doc.format == "docx"


def test_parse_jd_heuristic_must_and_nice() -> None:
    jd_text = (FIXTURES / "sample_jd.txt").read_text(encoding="utf-8")
    structured = parse_jd_structured(jd_text, use_llm=False)

    assert structured.domain == "technical"
    assert structured.seniority == "mid"
    assert structured.industry == "EXAai Demo Labs"
    assert any("Python" in item for item in structured.must_have)
    assert any("Kubernetes" in item for item in structured.nice_to_have)


def test_parse_jd_title() -> None:
    structured = parse_jd_structured("Product Designer\n\nRequirements:\n- Figma", use_llm=False)
    assert structured.job_title == "Product Designer"
    assert structured.domain == "design"


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        ("resume.pdf", "pdf"),
        ("resume.docx", "docx"),
        ("resume.txt", "txt"),
    ],
)
def test_detect_format_by_extension(filename: str, expected: str) -> None:
    assert detect_format(b"", filename) == expected
